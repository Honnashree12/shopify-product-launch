import os
import csv
import json
import logging
from typing import List, Dict, Any, Tuple
import pandas as pd
import pypdf
import docx
from PIL import Image
import cv2

logger = logging.getLogger("file_extractor")

SUPPORTED_IMAGE_EXTS = (".jpg", ".jpeg", ".png", ".webp", ".gif", ".heic")
SUPPORTED_VIDEO_EXTS = (".mp4", ".mov")

def scan_directory(dir_path: str) -> List[str]:
    """Recursively lists all files in a directory."""
    file_list = []
    for root, _, files in os.walk(dir_path):
        for f in files:
            file_list.append(os.path.abspath(os.path.join(root, f)))
    return file_list

def extract_pdf(file_path: str) -> str:
    """Extracts text content and attempts to extract embedded images from a PDF."""
    try:
        reader = pypdf.PdfReader(file_path)
        text_content = []
        for page_num, page in enumerate(reader.pages):
            text = page.extract_text()
            if text:
                text_content.append(text)
            
            # Extract images from page if technically possible
            try:
                for img_idx, image_file_object in enumerate(page.images):
                    img_ext = os.path.splitext(image_file_object.name)[1].lower() or ".png"
                    img_name = f"extracted_pdf_{os.path.basename(file_path)}_{page_num}_{img_idx}{img_ext}"
                    img_dest = os.path.join(os.path.dirname(file_path), img_name)
                    with open(img_dest, "wb") as fp:
                        fp.write(image_file_object.data)
                    logger.info("Extracted embedded image from PDF: %s", img_dest)
            except Exception as img_err:
                logger.debug("Failed to extract page images from PDF: %s", img_err)

        return "\n".join(text_content)
    except Exception as e:
        logger.error("Failed to parse PDF %s: %s", file_path, e)
        return f"[PDF Parsing Error for {os.path.basename(file_path)}: {e}]"

def extract_docx(file_path: str) -> str:
    """Extracts text content from a Word Document (.docx), including tables."""
    try:
        doc = docx.Document(file_path)
        text_content = []
        for p in doc.paragraphs:
            if p.text.strip():
                text_content.append(p.text.strip())
        
        # Extract tables as markdown-like rows
        for t in doc.tables:
            for row in t.rows:
                cells = [cell.text.strip() for cell in row.cells]
                if any(cells):
                    text_content.append(" | ".join(cells))
                    
        return "\n".join(text_content)
    except Exception as e:
        logger.error("Failed to parse DOCX %s: %s", file_path, e)
        return f"[DOCX Parsing Error for {os.path.basename(file_path)}: {e}]"

def extract_xlsx(file_path: str) -> str:
    """Extracts data from Excel spreadsheets using pandas."""
    try:
        xls = pd.ExcelFile(file_path)
        lines = []
        for sheet_name in xls.sheet_names:
            df = xls.parse(sheet_name)
            lines.append(f"### Sheet: {sheet_name} ###")
            lines.append(df.to_string(index=False))
        return "\n".join(lines)
    except Exception as e:
        logger.error("Failed to parse XLSX %s: %s", file_path, e)
        return f"[Excel Parsing Error for {os.path.basename(file_path)}: {e}]"

def extract_csv(file_path: str) -> str:
    """Extracts data from CSV file."""
    try:
        lines = []
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            reader = csv.reader(f)
            for row in reader:
                if any(row):
                    lines.append(", ".join(row))
        return "\n".join(lines)
    except Exception as e:
        logger.error("Failed to parse CSV %s: %s", file_path, e)
        return f"[CSV Parsing Error for {os.path.basename(file_path)}: {e}]"

def extract_json(file_path: str) -> str:
    """Reads JSON file and returns structured text representation."""
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            data = json.load(f)
        return json.dumps(data, indent=2)
    except Exception as e:
        logger.error("Failed to parse JSON %s: %s", file_path, e)
        return f"[JSON Parsing Error for {os.path.basename(file_path)}: {e}]"

def get_image_metadata(file_path: str) -> Dict[str, Any]:
    """Retrieves dimensions, size, and type information for an image."""
    try:
        with Image.open(file_path) as img:
            width, height = img.size
            mime = getattr(img, "get_format_mimetype", lambda: "image/png")()
            return {
                "filename": os.path.basename(file_path),
                "file_path": file_path,
                "extension": os.path.splitext(file_path)[1].lower(),
                "mime_type": mime,
                "dimensions": f"{width}x{height}",
                "file_size": os.path.getsize(file_path)
            }
    except Exception as e:
        logger.error("Failed to read image metadata for %s: %s", file_path, e)
        return {
            "filename": os.path.basename(file_path),
            "file_path": file_path,
            "extension": os.path.splitext(file_path)[1].lower(),
            "error": str(e),
            "file_size": os.path.getsize(file_path)
        }

def get_video_metadata(file_path: str) -> Dict[str, Any]:
    """Retrieves duration, dimensions, size, and codec information for a video using OpenCV."""
    try:
        cap = cv2.VideoCapture(file_path)
        if not cap.isOpened():
            raise ValueError("Failed to open video file stream.")
        
        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        fps = cap.get(cv2.CAP_PROP_FPS)
        frame_count = int(cap.get(cv2.CAP_PROP_FRAME_COUNT))
        duration = frame_count / fps if fps > 0 else 0
        
        fourcc = int(cap.get(cv2.CAP_PROP_FOURCC))
        codec = "".join([chr((fourcc >> 8 * i) & 0xFF) for i in range(4)]).strip()
        cap.release()
        
        ext = os.path.splitext(file_path)[1].lower()
        mime_type = "video/mp4" if ext == ".mp4" else ("video/quicktime" if ext == ".mov" else "video/webm")
        
        return {
            "filename": os.path.basename(file_path),
            "file_path": file_path,
            "extension": ext,
            "mime_type": mime_type,
            "dimensions": f"{width}x{height}",
            "duration": round(duration, 2),
            "codec": codec or "unknown",
            "file_size": os.path.getsize(file_path)
        }
    except Exception as e:
        logger.error("Failed to read video metadata for %s: %s", file_path, e)
        return {
            "filename": os.path.basename(file_path),
            "file_path": file_path,
            "extension": os.path.splitext(file_path)[1].lower(),
            "error": str(e),
            "file_size": os.path.getsize(file_path)
        }

def convert_image_if_needed(file_path: str) -> Tuple[str, bool]:
    """Converts unsupported image format (e.g. TIFF, BMP) to PNG, keeping original intact."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext in SUPPORTED_IMAGE_EXTS:
        return file_path, False
    
    # Try safe conversion using Pillow
    try:
        out_path = os.path.splitext(file_path)[0] + "_converted.png"
        with Image.open(file_path) as img:
            img.save(out_path, "PNG")
        logger.info("Converted unsupported image %s to PNG: %s", file_path, out_path)
        return out_path, True
    except Exception as e:
        logger.error("Failed to convert image %s to PNG: %s", file_path, e)
        return file_path, False

def convert_video_if_needed(file_path: str) -> Tuple[str, bool]:
    """Converts unsupported video formats (e.g. AVI, MKV) to MP4 frame-by-frame, keeping original intact."""
    ext = os.path.splitext(file_path)[1].lower()
    if ext in SUPPORTED_VIDEO_EXTS:
        return file_path, False
    
    # Try safe frame-by-frame conversion using cv2 VideoWriter
    try:
        out_path = os.path.splitext(file_path)[0] + "_converted.mp4"
        cap = cv2.VideoCapture(file_path)
        if not cap.isOpened():
            raise ValueError("Failed to open source video file.")
            
        width = int(cap.get(cv2.CAP_PROP_FRAME_WIDTH))
        height = int(cap.get(cv2.CAP_PROP_FRAME_HEIGHT))
        fps = cap.get(cv2.CAP_PROP_FPS) or 24.0
        
        fourcc = cv2.VideoWriter_fourcc(*'mp4v')
        out = cv2.VideoWriter(out_path, fourcc, fps, (width, height))
        
        while True:
            ret, frame = cap.read()
            if not ret:
                break
            out.write(frame)
            
        cap.release()
        out.release()
        
        logger.info("Converted unsupported video %s to MP4: %s", file_path, out_path)
        return out_path, True
    except Exception as e:
        logger.error("Failed to convert video %s to MP4: %s", file_path, e)
        return file_path, False
